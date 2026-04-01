import io
from typing import List, Optional, Tuple
from fastapi import UploadFile
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func
from app.models.document import Document, DocumentStatus
from app.models.user import User
from app.services.storage_service import storage_service
from app.services.embedding_service import embedding_service
from app.core.config import settings
from app.core.exceptions import ValidationError, NotFoundError, AuthorizationError
from app.core.logging import get_logger

logger = get_logger(__name__)


def extract_text_from_pdf(content: bytes) -> str:
    import PyPDF2
    reader = PyPDF2.PdfReader(io.BytesIO(content))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def extract_text_from_docx(content: bytes) -> str:
    import docx
    doc = docx.Document(io.BytesIO(content))
    return "\n".join(para.text for para in doc.paragraphs if para.text.strip())


def extract_text_from_txt(content: bytes) -> str:
    import chardet
    detected = chardet.detect(content)
    encoding = detected.get("encoding", "utf-8") or "utf-8"
    return content.decode(encoding, errors="replace")


def extract_text(content: bytes, content_type: str) -> str:
    if content_type == "application/pdf":
        return extract_text_from_pdf(content)
    elif "wordprocessingml" in content_type:
        return extract_text_from_docx(content)
    elif content_type == "text/plain":
        return extract_text_from_txt(content)
    raise ValueError(f"Unsupported file type: {content_type}")


class DocumentService:
    def __init__(self, db: AsyncSession):
        self.db = db

    async def upload(
        self,
        file: UploadFile,
        user: User,
        project_id: Optional[str] = None,
    ) -> Document:
        if file.content_type not in settings.allowed_file_types:
            raise ValidationError(
                f"File type '{file.content_type}' not allowed. "
                f"Allowed: {settings.allowed_file_types}"
            )

        content = await file.read()
        if len(content) > settings.max_file_size_bytes:
            raise ValidationError(f"File too large. Maximum: {settings.max_file_size_mb}MB")

        s3_key, s3_bucket = await storage_service.upload_file(
            file_content=content,
            filename=file.filename,
            content_type=file.content_type,
            user_id=user.id,
        )

        document = Document(
            filename=s3_key.split("/")[-1],
            original_filename=file.filename,
            file_type=file.content_type,
            file_size=len(content),
            s3_key=s3_key,
            s3_bucket=s3_bucket,
            project_id=project_id,
            uploaded_by=user.id,
            status=DocumentStatus.PROCESSING.value,
        )
        self.db.add(document)
        await self.db.flush()
        await self.db.refresh(document)

        try:
            extracted_text = extract_text(content, file.content_type)
            chunk_count = await embedding_service.add_document(
                document_id=document.id,
                text=extracted_text,
            )
            document.extracted_text = extracted_text
            document.chunk_count = chunk_count
            document.status = DocumentStatus.PROCESSED.value
        except Exception as e:
            document.status = DocumentStatus.FAILED.value
            document.error_message = str(e)
            logger.error("document_processing_failed", doc_id=document.id, error=str(e))

        await self.db.flush()
        await self.db.refresh(document)
        return document

    async def get_by_id(self, document_id: str, user: User) -> Document:
        result = await self.db.execute(select(Document).where(Document.id == document_id))
        document = result.scalar_one_or_none()
        if not document:
            raise NotFoundError("Document", document_id)
        if document.uploaded_by != user.id:
            raise AuthorizationError("You do not have access to this document")
        return document

    async def list_for_user(
        self,
        user: User,
        project_id: Optional[str] = None,
        page: int = 1,
        page_size: int = 20,
    ) -> Tuple[List[Document], int]:
        query = select(Document).where(Document.uploaded_by == user.id)
        count_query = select(func.count(Document.id)).where(Document.uploaded_by == user.id)

        if project_id:
            query = query.where(Document.project_id == project_id)
            count_query = count_query.where(Document.project_id == project_id)

        total = (await self.db.execute(count_query)).scalar_one()
        offset = (page - 1) * page_size
        documents = (
            await self.db.execute(
                query.order_by(Document.created_at.desc()).offset(offset).limit(page_size)
            )
        ).scalars().all()

        return list(documents), total

    async def delete(self, document: Document, user: User) -> None:
        if document.uploaded_by != user.id:
            raise AuthorizationError("You can only delete your own documents")
        await storage_service.delete_file(document.s3_key)
        await embedding_service.remove_document(document.id)
        await self.db.delete(document)
        await self.db.flush()
        logger.info("document_deleted", doc_id=document.id)

    async def get_total_count(self) -> int:
        result = await self.db.execute(select(func.count(Document.id)))
        return result.scalar_one()

    async def get_presigned_url(self, document: Document) -> str:
        return await storage_service.get_presigned_url(document.s3_key)
